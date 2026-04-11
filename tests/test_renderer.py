"""
Regression tests for ATS-compliant DOCX renderer.

These tests verify that the renderer produces documents that
will parse correctly on Applicant Tracking Systems.
"""

import json
from pathlib import Path

import pytest
from docx import Document

from jobtool.models import MasterCV
from jobtool.renderer.docx_renderer import render_cv, render_cover_letter


# Path to test fixture
FIXTURE_PATH = Path(__file__).parent / "fixtures" / "sample-master-cv.json"


@pytest.fixture
def sample_cv() -> MasterCV:
    """Load the sample Master CV fixture."""
    with open(FIXTURE_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    return MasterCV.model_validate(data)


@pytest.fixture
def output_dir(tmp_path: Path) -> Path:
    """Create a temporary output directory."""
    return tmp_path / "output"


class TestATSCompliance:
    """Test that rendered documents meet ATS compliance requirements."""

    def test_no_tables(self, sample_cv: MasterCV, output_dir: Path):
        """CV must not contain any tables."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        assert len(doc.tables) == 0, "ATS-compliant CV must not contain tables"

    def test_no_headers(self, sample_cv: MasterCV, output_dir: Path):
        """CV must not use document headers."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        for section in doc.sections:
            header = section.header
            # Header should be empty or contain no text
            header_text = "".join(p.text for p in header.paragraphs)
            assert header_text.strip() == "", "ATS-compliant CV must not use headers"

    def test_no_footers(self, sample_cv: MasterCV, output_dir: Path):
        """CV must not use document footers."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        for section in doc.sections:
            footer = section.footer
            # Footer should be empty or contain no text
            footer_text = "".join(p.text for p in footer.paragraphs)
            assert footer_text.strip() == "", "ATS-compliant CV must not use footers"

    def test_uses_arial_font(self, sample_cv: MasterCV, output_dir: Path):
        """CV must use Arial font throughout."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        for para in doc.paragraphs:
            for run in para.runs:
                # Font name should be Arial (or None, which inherits from style)
                if run.font.name:
                    assert run.font.name == "Arial", f"Found non-Arial font: {run.font.name}"

    def test_has_required_sections(self, sample_cv: MasterCV, output_dir: Path):
        """CV must have all required section headings."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        full_text = "\n".join(p.text for p in doc.paragraphs)

        required_sections = [
            "Personal Statement",
            "Work Experience",
            "Education",
            "Skills",
            "References",
        ]

        for section in required_sections:
            assert section in full_text, f"Missing required section: {section}"

    def test_ends_with_references(self, sample_cv: MasterCV, output_dir: Path):
        """CV must end with 'References available on request'."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        # Get all non-empty paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Last paragraph should be the references line
        assert "Available on request" in paragraphs[-1], \
            "CV must end with 'References available on request'"

    def test_contact_in_body(self, sample_cv: MasterCV, output_dir: Path):
        """Contact details must be in the document body, not header."""
        output_file = render_cv(sample_cv, output_dir, "Test Job")
        doc = Document(str(output_file))

        body_text = "\n".join(p.text for p in doc.paragraphs)

        # Contact details should appear in body
        assert sample_cv.personalDetails.email in body_text or "REPLACE_WITH" in sample_cv.personalDetails.email, \
            "Contact email must be in document body"

    def test_filename_format(self, sample_cv: MasterCV, output_dir: Path):
        """Filename must follow the correct format."""
        output_file = render_cv(sample_cv, output_dir, "Data Entry Clerk")

        # Filename should be FirstName-LastName-JobTitle-Year.docx
        filename = output_file.name
        assert filename.startswith("John-Smith-"), f"Unexpected filename: {filename}"
        assert "data-entry-clerk" in filename.lower(), f"Job title missing from filename: {filename}"
        assert filename.endswith(".docx"), f"Wrong extension: {filename}"


class TestCoverLetter:
    """Test cover letter rendering."""

    def test_cover_letter_renders(self, output_dir: Path):
        """Cover letter should render without errors."""
        sample_text = """Dear Hiring Manager,

I am writing to apply for the Data Entry Clerk position at Acme Corporation.

With 5 years of experience in data entry and administration, I am confident
in my ability to contribute to your team.

I look forward to hearing from you.

Yours sincerely,
John Smith"""

        output_file = render_cover_letter(
            sample_text,
            output_dir,
            "John Smith",
            "Data Entry Clerk"
        )

        assert output_file.exists()
        assert output_file.suffix == ".docx"

    def test_cover_letter_uses_arial(self, output_dir: Path):
        """Cover letter must use Arial font."""
        sample_text = "This is a test cover letter."

        output_file = render_cover_letter(
            sample_text,
            output_dir,
            "John Smith",
            "Test Job"
        )

        doc = Document(str(output_file))

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    assert run.font.name == "Arial"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
