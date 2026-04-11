"""
ATS-Compliant DOCX Renderer

This is the most critical module in JobAutoApply. It generates DOCX files
that are guaranteed to parse correctly on all major Applicant Tracking Systems.

ATS COMPLIANCE RULES ENFORCED:
1. Single column only - NO Tables, NO text boxes, NO shapes
2. Arial font exclusively (11pt body, 14pt headings, 18pt name)
3. All content in document body - NO headers, NO footers
4. Exact section labels: Personal Statement, Work Experience, Education,
   Skills, Certifications, Languages, References
5. Solid round bullets only (bullet character)
6. 2cm margins on all sides
7. UK date format: Mon YYYY - Mon YYYY
8. British English spelling (enforced by AI, not renderer)
9. Black text only, no underlines
10. Filename: FirstName-LastName-JobTitleSlug-2026.docx
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from jobtool.models import TailoredCV, MasterCV, PersonalDetails


# ============================================================================
# Constants
# ============================================================================

FONT_NAME = "Arial"
FONT_SIZE_NAME = Pt(18)
FONT_SIZE_SECTION = Pt(14)
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_CONTACT = Pt(11)

MARGIN = Cm(2)

# Standard section headings (UK conventions)
SECTION_PERSONAL_STATEMENT = "Personal Statement"
SECTION_WORK_EXPERIENCE = "Work Experience"
SECTION_EDUCATION = "Education"
SECTION_SKILLS = "Skills"
SECTION_CERTIFICATIONS = "Certifications"
SECTION_LANGUAGES = "Languages"
SECTION_REFERENCES = "References"

# Month name mapping for date formatting
MONTH_NAMES = {
    "01": "Jan", "02": "Feb", "03": "Mar", "04": "Apr",
    "05": "May", "06": "Jun", "07": "Jul", "08": "Aug",
    "09": "Sep", "10": "Oct", "11": "Nov", "12": "Dec",
}


# ============================================================================
# Date Formatting
# ============================================================================

def format_date(date_str: str) -> str:
    """
    Format a date string from YYYY-MM to Mon YYYY format.

    Examples:
        "2024-01" -> "Jan 2024"
        "Present" -> "Present"
    """
    if not date_str or date_str.lower() == "present":
        return "Present"

    parts = date_str.split("-")
    if len(parts) != 2:
        return date_str

    year, month = parts
    month_name = MONTH_NAMES.get(month, month)
    return f"{month_name} {year}"


def format_date_range(start: str, end: str) -> str:
    """Format a date range as 'Mon YYYY - Mon YYYY'."""
    return f"{format_date(start)} - {format_date(end)}"


# ============================================================================
# Filename Generation
# ============================================================================

def slugify(text: str) -> str:
    """Convert text to a URL/filename-safe slug."""
    # Convert to lowercase and replace spaces with hyphens
    slug = text.lower().strip()
    slug = re.sub(r'[^\w\s-]', '', slug)  # Remove non-word chars except hyphens
    slug = re.sub(r'[-\s]+', '-', slug)   # Replace spaces/multiple hyphens
    return slug.strip('-')


def generate_cv_filename(personal_details: PersonalDetails, job_title: str) -> str:
    """
    Generate the CV filename per spec.

    Format: FirstName-LastName-JobTitleSlug-2026.docx
    """
    name_parts = personal_details.fullName.split()
    first_name = name_parts[0] if name_parts else "Unknown"
    last_name = name_parts[-1] if len(name_parts) > 1 else "User"

    title_slug = slugify(job_title)
    year = datetime.now().year

    return f"{first_name}-{last_name}-{title_slug}-{year}.docx"


# ============================================================================
# Document Setup
# ============================================================================

def _set_document_margins(doc: Document) -> None:
    """Set 2cm margins on all sides."""
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN


def _setup_styles(doc: Document) -> None:
    """
    Set up document styles for consistent formatting.

    We define custom styles rather than relying on built-in ones
    to ensure ATS compatibility across all systems.
    """
    styles = doc.styles

    # Name style (18pt Arial Bold)
    if "CV Name" not in [s.name for s in styles]:
        name_style = styles.add_style("CV Name", WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = FONT_NAME
        name_style.font.size = FONT_SIZE_NAME
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(0, 0, 0)
        name_style.paragraph_format.space_after = Pt(6)
        # Set font for complex scripts (required for some systems)
        name_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Section heading style (14pt Arial Bold)
    if "CV Section" not in [s.name for s in styles]:
        section_style = styles.add_style("CV Section", WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = FONT_NAME
        section_style.font.size = FONT_SIZE_SECTION
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0, 0, 0)
        section_style.paragraph_format.space_before = Pt(12)
        section_style.paragraph_format.space_after = Pt(6)
        section_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Body text style (11pt Arial)
    if "CV Body" not in [s.name for s in styles]:
        body_style = styles.add_style("CV Body", WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = FONT_NAME
        body_style.font.size = FONT_SIZE_BODY
        body_style.font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Bullet style (11pt Arial with bullet)
    if "CV Bullet" not in [s.name for s in styles]:
        bullet_style = styles.add_style("CV Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = FONT_NAME
        bullet_style.font.size = FONT_SIZE_BODY
        bullet_style.font.color.rgb = RGBColor(0, 0, 0)
        bullet_style.paragraph_format.space_after = Pt(3)
        bullet_style.paragraph_format.left_indent = Cm(0.5)
        bullet_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # Job title style (11pt Arial Bold)
    if "CV Job Title" not in [s.name for s in styles]:
        job_style = styles.add_style("CV Job Title", WD_STYLE_TYPE.PARAGRAPH)
        job_style.font.name = FONT_NAME
        job_style.font.size = FONT_SIZE_BODY
        job_style.font.bold = True
        job_style.font.color.rgb = RGBColor(0, 0, 0)
        job_style.paragraph_format.space_before = Pt(6)
        job_style.paragraph_format.space_after = Pt(3)
        job_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _ensure_font(paragraph, font_name: str = FONT_NAME) -> None:
    """Ensure the paragraph uses the specified font throughout."""
    for run in paragraph.runs:
        run.font.name = font_name
        # Set font for East Asian text (required for full compatibility)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)


# ============================================================================
# Content Rendering
# ============================================================================

def _add_name(doc: Document, name: str) -> None:
    """Add the candidate's name (18pt bold, centered or left-aligned)."""
    para = doc.add_paragraph(name, style="CV Name")
    _ensure_font(para)


def _add_contact_line(doc: Document, details: PersonalDetails) -> None:
    """
    Add contact details on a single line in the document body.

    Format: City, Region | phone | email | linkedin (if present)

    IMPORTANT: Contact details go in the BODY, not in headers/footers.
    """
    parts = []

    # Location
    location = f"{details.city}, {details.region}"
    parts.append(location)

    # Phone (if not placeholder)
    if details.phone and "REPLACE_WITH" not in details.phone:
        parts.append(details.phone)

    # Email (if not placeholder)
    if details.email and "REPLACE_WITH" not in details.email:
        parts.append(details.email)

    # LinkedIn (if present and not placeholder)
    if details.linkedin and "REPLACE_WITH" not in details.linkedin:
        # Clean up LinkedIn URL to just show the profile path
        linkedin = details.linkedin.replace("https://", "").replace("www.", "")
        parts.append(linkedin)

    contact_text = " | ".join(parts)
    para = doc.add_paragraph(contact_text, style="CV Body")
    para.paragraph_format.space_after = Pt(12)
    _ensure_font(para)


def _add_section_heading(doc: Document, heading: str) -> None:
    """Add a section heading (14pt bold)."""
    para = doc.add_paragraph(heading, style="CV Section")
    _ensure_font(para)


def _add_body_text(doc: Document, text: str) -> None:
    """Add body text (11pt)."""
    para = doc.add_paragraph(text, style="CV Body")
    _ensure_font(para)


def _add_bullet(doc: Document, text: str) -> None:
    """
    Add a bullet point.

    Uses a solid round bullet character (Unicode 2022) followed by text.
    We manually add the bullet rather than using Word's list styles
    for maximum ATS compatibility.
    """
    # Use bullet character directly
    bullet_text = f"\u2022 {text}"
    para = doc.add_paragraph(bullet_text, style="CV Bullet")
    _ensure_font(para)


def _add_job_entry(doc: Document, job_title: str, employer: str,
                   location: str, date_range: str) -> None:
    """
    Add a work experience entry header.

    Format:
    Job Title | Employer | Location
    Mon YYYY - Mon YYYY
    """
    # Title line
    title_line = f"{job_title} | {employer} | {location}"
    para = doc.add_paragraph(title_line, style="CV Job Title")
    _ensure_font(para)

    # Date line
    date_para = doc.add_paragraph(date_range, style="CV Body")
    date_para.paragraph_format.space_after = Pt(3)
    _ensure_font(date_para)


def _add_education_entry(doc: Document, degree: str, institution: str,
                         location: str, date_range: str, grade: str | None) -> None:
    """Add an education entry."""
    # Degree and institution
    title_line = f"{degree} | {institution} | {location}"
    para = doc.add_paragraph(title_line, style="CV Job Title")
    _ensure_font(para)

    # Date and grade
    date_text = date_range
    if grade and "REPLACE_WITH" not in grade:
        date_text += f" | {grade}"

    date_para = doc.add_paragraph(date_text, style="CV Body")
    date_para.paragraph_format.space_after = Pt(6)
    _ensure_font(date_para)


# ============================================================================
# Main Render Functions
# ============================================================================

def render_cv(cv: TailoredCV | MasterCV, output_path: Path, job_title: str = "CV") -> Path:
    """
    Render a CV to an ATS-compliant DOCX file.

    Args:
        cv: The tailored CV data (TailoredCV or MasterCV)
        output_path: Directory to save the file, or full file path
        job_title: Job title for filename generation (if output_path is a directory)

    Returns:
        Path to the generated DOCX file

    ATS Compliance:
        - Single column layout (no Tables)
        - Arial font throughout
        - All content in document body
        - Standard UK section headings
        - 2cm margins
        - Solid round bullets
    """
    doc = Document()

    # Setup document
    _set_document_margins(doc)
    _setup_styles(doc)

    # Determine output file path
    # If output_path doesn't have a .docx extension, treat it as a directory
    if output_path.suffix.lower() != ".docx":
        # It's a directory (existing or to be created)
        output_path.mkdir(parents=True, exist_ok=True)
        filename = generate_cv_filename(cv.personalDetails, job_title)
        output_file = output_path / filename
    else:
        # It's a full file path
        output_file = output_path
        output_file.parent.mkdir(parents=True, exist_ok=True)

    # ========== NAME ==========
    _add_name(doc, cv.personalDetails.fullName)

    # ========== CONTACT (in body, NOT header) ==========
    _add_contact_line(doc, cv.personalDetails)

    # ========== PERSONAL STATEMENT ==========
    _add_section_heading(doc, SECTION_PERSONAL_STATEMENT)
    _add_body_text(doc, cv.personalStatement)

    # ========== WORK EXPERIENCE ==========
    if cv.workExperience:
        _add_section_heading(doc, SECTION_WORK_EXPERIENCE)

        for exp in cv.workExperience:
            date_range = format_date_range(exp.startDate, exp.endDate)
            _add_job_entry(doc, exp.jobTitle, exp.employer, exp.location, date_range)

            for bullet in exp.bullets:
                _add_bullet(doc, bullet)

    # ========== EDUCATION ==========
    if cv.education:
        _add_section_heading(doc, SECTION_EDUCATION)

        for edu in cv.education:
            date_range = format_date_range(edu.startDate, edu.endDate)
            _add_education_entry(
                doc, edu.degree, edu.institution, edu.location,
                date_range, edu.grade
            )

            for highlight in edu.highlights:
                _add_bullet(doc, highlight)

    # ========== SKILLS ==========
    if cv.skills:
        _add_section_heading(doc, SECTION_SKILLS)

        # Combine all skills into categorized lines
        skill_lines = []

        if cv.skills.technical:
            skill_lines.append(f"Technical: {', '.join(cv.skills.technical)}")

        if cv.skills.soft:
            skill_lines.append(f"Soft Skills: {', '.join(cv.skills.soft)}")

        if cv.skills.tools:
            skill_lines.append(f"Tools: {', '.join(cv.skills.tools)}")

        for line in skill_lines:
            _add_body_text(doc, line)

    # ========== CERTIFICATIONS ==========
    if cv.certifications:
        _add_section_heading(doc, SECTION_CERTIFICATIONS)

        for cert in cv.certifications:
            # Skip if key fields are placeholders
            if "REPLACE_WITH" in cert.issueDate:
                continue

            cert_text = f"{cert.name} - {cert.issuer} ({format_date(cert.issueDate)})"
            _add_body_text(doc, cert_text)

    # ========== LANGUAGES ==========
    if cv.languages:
        _add_section_heading(doc, SECTION_LANGUAGES)

        lang_parts = [f"{lang.language} ({lang.proficiency})" for lang in cv.languages]
        _add_body_text(doc, ", ".join(lang_parts))

    # ========== REFERENCES ==========
    _add_section_heading(doc, SECTION_REFERENCES)
    _add_body_text(doc, cv.references)

    # Save the document
    doc.save(str(output_file))

    return output_file


def render_cover_letter(
    text: str,
    output_path: Path,
    candidate_name: str,
    job_title: str = "Cover-Letter"
) -> Path:
    """
    Render a cover letter to a DOCX file.

    Args:
        text: The cover letter text (plain text)
        output_path: Directory to save the file, or full file path
        candidate_name: Candidate's full name for filename
        job_title: Job title for filename

    Returns:
        Path to the generated DOCX file
    """
    doc = Document()

    # Setup document
    _set_document_margins(doc)
    _setup_styles(doc)

    # Determine output file path
    # If output_path doesn't have a .docx extension, treat it as a directory
    if output_path.suffix.lower() != ".docx":
        # It's a directory (existing or to be created)
        output_path.mkdir(parents=True, exist_ok=True)
        name_parts = candidate_name.split()
        first_name = name_parts[0] if name_parts else "Unknown"
        last_name = name_parts[-1] if len(name_parts) > 1 else "User"
        title_slug = slugify(job_title)
        year = datetime.now().year
        filename = f"{first_name}-{last_name}-{title_slug}-Cover-Letter-{year}.docx"
        output_file = output_path / filename
    else:
        # It's a full file path
        output_file = output_path
        output_file.parent.mkdir(parents=True, exist_ok=True)

    # Add date
    today = datetime.now().strftime("%d %B %Y")
    date_para = doc.add_paragraph(today, style="CV Body")
    date_para.paragraph_format.space_after = Pt(12)
    _ensure_font(date_para)

    # Add cover letter body
    # Split by double newlines to preserve paragraph structure
    paragraphs = text.strip().split("\n\n")

    for para_text in paragraphs:
        # Clean up single newlines within paragraphs
        clean_text = para_text.replace("\n", " ").strip()
        if clean_text:
            para = doc.add_paragraph(clean_text, style="CV Body")
            para.paragraph_format.space_after = Pt(12)
            _ensure_font(para)

    # Save the document
    doc.save(str(output_file))

    return output_file
